from __future__ import annotations

import os
import re
import shutil
import sys
from pathlib import Path

WORKSPACE = Path(__file__).resolve().parents[1]
DEPS = WORKSPACE / ".codex_deps"
if DEPS.exists():
    sys.path.insert(0, str(DEPS))

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import (
    WD_ALIGN_PARAGRAPH,
    WD_BREAK,
    WD_LINE_SPACING,
    WD_TAB_ALIGNMENT,
    WD_TAB_LEADER,
)
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


SOURCE_MD = WORKSPACE / "docs" / "PROJECT_DEVELOPMENT_DOCUMENT.md"
COVER_SOURCE = Path(
    r"C:\Users\Administrator\.codex\generated_images\01a0529c-24c7-71e2-b400-957857d761e9\exec-bedc4de7-ec09-4128-ac49-67eee03a3ae8.png"
)
ASSET_DIR = WORKSPACE / "assets" / "document"
OUTPUT_DIR = WORKSPACE / "deliverables"
OUTPUT_DOCX = OUTPUT_DIR / "归音_项目开发文档_v1.0.docx"

# narrative_proposal preset, resolved as exact tokens.
PAGE_WIDTH_IN = 8.5
PAGE_HEIGHT_IN = 11.0
MARGIN_IN = 1.0
CONTENT_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120
FONT_LATIN = "Calibri"
FONT_EAST_ASIA = "Noto Sans SC"
FONT_DISPLAY = "Noto Serif SC"
FONT_MONO = "Consolas"

NAVY = "17324D"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
TEAL = "267C7A"
TEAL_LIGHT = "DDEFEA"
AMBER = "D9923B"
AMBER_LIGHT = "FFF3E0"
INK = "243746"
MUTED = "667784"
LIGHT = "F4F6F9"
MID = "E2E9EE"
WHITE = "FFFFFF"
RED = "9B1C1C"


def rgb(hex_value: str) -> RGBColor:
    return RGBColor.from_string(hex_value)


def set_run_font(run, name=FONT_LATIN, east_asia=FONT_EAST_ASIA, size=None,
                 color=None, bold=None, italic=None):
    run.font.name = name
    rpr = run._element.get_or_add_rPr()
    fonts = rpr.rFonts
    if fonts is None:
        fonts = OxmlElement("w:rFonts")
        rpr.insert(0, fonts)
    fonts.set(qn("w:ascii"), name)
    fonts.set(qn("w:hAnsi"), name)
    fonts.set(qn("w:eastAsia"), east_asia)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = rgb(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_repeat_table_header(row):
    trpr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    trpr.append(tbl_header)


def set_cant_split(row):
    trpr = row._tr.get_or_add_trPr()
    cant_split = trpr.find(qn("w:cantSplit"))
    if cant_split is None:
        cant_split = OxmlElement("w:cantSplit")
        trpr.append(cant_split)


def set_cell_shading(cell, fill):
    tcpr = cell._tc.get_or_add_tcPr()
    shd = tcpr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tcpr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tcpr = tc.get_or_add_tcPr()
    tcmar = tcpr.first_child_found_in("w:tcMar")
    if tcmar is None:
        tcmar = OxmlElement("w:tcMar")
        tcpr.append(tcmar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tcmar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tcmar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width_dxa):
    tcpr = cell._tc.get_or_add_tcPr()
    tcw = tcpr.find(qn("w:tcW"))
    if tcw is None:
        tcw = OxmlElement("w:tcW")
        tcpr.append(tcw)
    tcw.set(qn("w:w"), str(width_dxa))
    tcw.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    assert sum(widths) == CONTENT_WIDTH_DXA, widths
    tblpr = table._tbl.tblPr
    tblw = tblpr.find(qn("w:tblW"))
    if tblw is None:
        tblw = OxmlElement("w:tblW")
        tblpr.append(tblw)
    tblw.set(qn("w:w"), str(CONTENT_WIDTH_DXA))
    tblw.set(qn("w:type"), "dxa")
    tblind = tblpr.find(qn("w:tblInd"))
    if tblind is None:
        tblind = OxmlElement("w:tblInd")
        tblpr.append(tblind)
    tblind.set(qn("w:w"), str(TABLE_INDENT_DXA))
    tblind.set(qn("w:type"), "dxa")
    layout = tblpr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tblpr.append(layout)
    layout.set(qn("w:type"), "fixed")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            set_cell_width(cell, width)
            set_cell_margins(cell)


def set_table_borders(table, color="CBD6DE", size=5):
    tblpr = table._tbl.tblPr
    borders = tblpr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tblpr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = borders.find(qn(f"w:{edge}"))
        if tag is None:
            tag = OxmlElement(f"w:{edge}")
            borders.append(tag)
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), str(size))
        tag.set(qn("w:color"), color)


def set_paragraph_shading(paragraph, fill):
    ppr = paragraph._p.get_or_add_pPr()
    shd = ppr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        ppr.append(shd)
    shd.set(qn("w:fill"), fill)


def add_bottom_border(paragraph, color=TEAL, size=14, space=5):
    ppr = paragraph._p.get_or_add_pPr()
    pbdr = ppr.find(qn("w:pBdr"))
    if pbdr is None:
        pbdr = OxmlElement("w:pBdr")
        ppr.append(pbdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size))
    bottom.set(qn("w:space"), str(space))
    bottom.set(qn("w:color"), color)
    pbdr.append(bottom)


def set_keep_with_next(paragraph, value=True):
    paragraph.paragraph_format.keep_with_next = value


def set_alt_text(inline_shape, title, description):
    docpr = inline_shape._inline.docPr
    docpr.set("title", title)
    docpr.set("descr", description)


def add_field(paragraph, instruction, placeholder="1"):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = placeholder
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, text, end])
    set_run_font(run, size=9, color=MUTED)


def set_update_fields(document):
    settings = document.settings._element
    node = settings.find(qn("w:updateFields"))
    if node is None:
        node = OxmlElement("w:updateFields")
        settings.append(node)
    node.set(qn("w:val"), "true")


TOC_ENTRIES = [
    ("1. 文档目的", 6),
    ("2. 项目概述", 7),
    ("3. 用户与需求分析", 9),
    ("4. 产品形态与信息架构", 11),
    ("5. 功能需求", 13),
    ("6. 非功能需求", 21),
    ("7. 系统总体架构", 25),
    ("8. 实时语音交互设计", 28),
    ("9. AI 编排与提示词策略", 30),
    ("10. 数据模型", 32),
    ("11. API 初步设计", 38),
    ("12. 权限模型", 41),
    ("13. 隐私、合规与伦理要求", 42),
    ("14. 安全威胁模型", 44),
    ("15. 报告生成算法", 45),
    ("16. UI/UX 关键流程", 47),
    ("17. 测试策略", 48),
    ("18. MVP 范围与验收标准", 50),
    ("19. 研发计划", 53),
    ("20. 团队配置", 55),
    ("21. 成本模型", 55),
    ("22. 商业模式与试点路径", 56),
    ("23. 创业比赛展示方案", 57),
    ("24. 风险登记表", 59),
    ("25. 开发仓库建议结构", 61),
    ("26. 环境规划", 61),
    ("27. 上线检查清单", 62),
    ("28. 待确认决策", 63),
    ("29. 下一步执行顺序", 64),
    ("30. 项目成功定义", 64),
]


def add_toc_entries(document, entries):
    for title, page in entries:
        p = document.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(5)
        p.paragraph_format.line_spacing = 1.08
        p.paragraph_format.keep_together = True
        p.paragraph_format.tab_stops.add_tab_stop(
            Inches(6.15), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS
        )
        r = p.add_run(title)
        set_run_font(r, size=10.2, color=INK, bold=title.startswith(("7.", "12.", "16.", "19.", "25.")))
        r = p.add_run(f"\t{page}")
        set_run_font(r, size=9.8, color=MUTED)


def add_hyperlink(paragraph, text, url):
    part = paragraph.part
    rid = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), rid)
    new_run = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), BLUE)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    rpr.extend([color, underline])
    new_run.append(rpr)
    t = OxmlElement("w:t")
    t.text = text
    new_run.append(t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


INLINE_PATTERN = re.compile(r"(\*\*.+?\*\*|`.+?`|\[[^\]]+\]\([^)]+\))")


def add_inline(paragraph, text, size=11, color=INK):
    pos = 0
    for match in INLINE_PATTERN.finditer(text):
        if match.start() > pos:
            run = paragraph.add_run(text[pos:match.start()])
            set_run_font(run, size=size, color=color)
        token = match.group(0)
        if token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            set_run_font(run, size=size, color=color, bold=True)
        elif token.startswith("`"):
            run = paragraph.add_run(token[1:-1])
            set_run_font(run, name=FONT_MONO, east_asia=FONT_EAST_ASIA, size=max(8.5, size - 1), color=DARK_BLUE)
            set_paragraph_shading(paragraph, "F2F6F8")
        else:
            m = re.match(r"\[([^\]]+)\]\(([^)]+)\)", token)
            add_hyperlink(paragraph, m.group(1), m.group(2))
        pos = match.end()
    if pos < len(text):
        run = paragraph.add_run(text[pos:])
        set_run_font(run, size=size, color=color)


def add_body_paragraph(document, text):
    p = document.add_paragraph(style="Normal")
    add_inline(p, text)
    return p


def ensure_numbering(document):
    numbering = document.part.numbering_part.element
    existing_abstract = [int(x.get(qn("w:abstractNumId"))) for x in numbering.findall(qn("w:abstractNum"))]
    existing_num = [int(x.get(qn("w:numId"))) for x in numbering.findall(qn("w:num"))]
    next_abs = max(existing_abstract or [0]) + 1
    next_num = max(existing_num or [0]) + 1

    def build(fmt, text, font=None):
        nonlocal next_abs, next_num
        abstract_id = next_abs
        num_id = next_num
        next_abs += 1
        next_num += 1
        abstract = OxmlElement("w:abstractNum")
        abstract.set(qn("w:abstractNumId"), str(abstract_id))
        multi = OxmlElement("w:multiLevelType")
        multi.set(qn("w:val"), "singleLevel")
        abstract.append(multi)
        lvl = OxmlElement("w:lvl")
        lvl.set(qn("w:ilvl"), "0")
        start = OxmlElement("w:start")
        start.set(qn("w:val"), "1")
        numfmt = OxmlElement("w:numFmt")
        numfmt.set(qn("w:val"), fmt)
        lvltext = OxmlElement("w:lvlText")
        lvltext.set(qn("w:val"), text)
        suff = OxmlElement("w:suff")
        suff.set(qn("w:val"), "tab")
        ppr = OxmlElement("w:pPr")
        tabs = OxmlElement("w:tabs")
        tab = OxmlElement("w:tab")
        tab.set(qn("w:val"), "num")
        tab.set(qn("w:pos"), "540")
        tabs.append(tab)
        ind = OxmlElement("w:ind")
        ind.set(qn("w:left"), "540")
        ind.set(qn("w:hanging"), "280")
        spacing = OxmlElement("w:spacing")
        spacing.set(qn("w:after"), "80")
        spacing.set(qn("w:line"), "290")
        spacing.set(qn("w:lineRule"), "auto")
        ppr.extend([tabs, ind, spacing])
        lvl.extend([start, numfmt, lvltext, suff, ppr])
        if font:
            rpr = OxmlElement("w:rPr")
            rfonts = OxmlElement("w:rFonts")
            rfonts.set(qn("w:ascii"), font)
            rfonts.set(qn("w:hAnsi"), font)
            rpr.append(rfonts)
            lvl.append(rpr)
        abstract.append(lvl)
        numbering.append(abstract)
        num = OxmlElement("w:num")
        num.set(qn("w:numId"), str(num_id))
        absref = OxmlElement("w:abstractNumId")
        absref.set(qn("w:val"), str(abstract_id))
        num.append(absref)
        numbering.append(num)
        return num_id

    return build("bullet", "•", FONT_LATIN), build("decimal", "%1.")


def apply_num(paragraph, num_id):
    ppr = paragraph._p.get_or_add_pPr()
    numpr = ppr.find(qn("w:numPr"))
    if numpr is None:
        numpr = OxmlElement("w:numPr")
        ppr.append(numpr)
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    nid = OxmlElement("w:numId")
    nid.set(qn("w:val"), str(num_id))
    numpr.extend([ilvl, nid])


def create_number_instance(document, base_num_id, start_value):
    numbering = document.part.numbering_part.element
    base = next(x for x in numbering.findall(qn("w:num")) if x.get(qn("w:numId")) == str(base_num_id))
    abstract_id = base.find(qn("w:abstractNumId")).get(qn("w:val"))
    existing = [int(x.get(qn("w:numId"))) for x in numbering.findall(qn("w:num"))]
    num_id = max(existing or [0]) + 1
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    absref = OxmlElement("w:abstractNumId")
    absref.set(qn("w:val"), abstract_id)
    override = OxmlElement("w:lvlOverride")
    override.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:startOverride")
    start.set(qn("w:val"), str(start_value))
    override.append(start)
    num.extend([absref, override])
    numbering.append(num)
    return num_id


def choose_widths(rows):
    n = len(rows[0])
    weights = []
    for col in range(n):
        lengths = [max(2, min(40, len(re.sub(r"[*`\[\]]", "", row[col])))) for row in rows]
        score = max(lengths) + sum(lengths) / max(1, len(lengths))
        weights.append(max(7.0, score))
    raw = [CONTENT_WIDTH_DXA * w / sum(weights) for w in weights]
    min_width = 1050 if n >= 4 else 1350
    widths = [max(min_width, int(x)) for x in raw]
    while sum(widths) > CONTENT_WIDTH_DXA:
        idx = max(range(n), key=lambda i: widths[i] - min_width)
        widths[idx] -= min(20, sum(widths) - CONTENT_WIDTH_DXA)
    widths[-1] += CONTENT_WIDTH_DXA - sum(widths)
    return widths


def add_markdown_table(document, rows):
    if not rows:
        return
    cols = len(rows[0])
    table = document.add_table(rows=len(rows), cols=cols)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    widths = choose_widths(rows)
    set_table_geometry(table, widths)
    set_table_borders(table)
    for r_idx, row in enumerate(rows):
        set_cant_split(table.rows[r_idx])
        for c_idx, value in enumerate(row):
            cell = table.cell(r_idx, c_idx)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.12
            if r_idx == 0:
                p.paragraph_format.keep_with_next = True
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if (r_idx == 0 or len(value) < 8) else WD_ALIGN_PARAGRAPH.LEFT
            add_inline(p, value, size=9.2 if cols >= 4 else 9.6, color=NAVY if r_idx == 0 else INK)
            if r_idx == 0:
                for run in p.runs:
                    run.bold = True
                set_cell_shading(cell, "E8EEF5")
            elif r_idx % 2 == 0:
                set_cell_shading(cell, "F8FAFB")
    set_repeat_table_header(table.rows[0])
    spacer = document.add_paragraph()
    spacer.paragraph_format.space_after = Pt(2)


def add_callout(document, label, text, fill=TEAL_LIGHT, accent=TEAL):
    table = document.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    set_table_geometry(table, [CONTENT_WIDTH_DXA])
    set_table_borders(table, color=accent, size=10)
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    set_cell_margins(cell, top=150, bottom=150, start=TABLE_INDENT_DXA, end=TABLE_INDENT_DXA)
    set_repeat_table_header(table.rows[0])
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    lead = p.add_run(label + "  ")
    set_run_font(lead, size=10.5, color=accent, bold=True)
    add_inline(p, text, size=10.5, color=INK)
    document.add_paragraph().paragraph_format.space_after = Pt(1)


def font(size, bold=False):
    path = Path(r"C:\Windows\Fonts\NotoSansSC-VF.ttf")
    if bold:
        path = Path(r"C:\Windows\Fonts\NotoSansSC-VF.ttf")
    return ImageFont.truetype(str(path), size=size)


def centered_text(draw, box, text, fnt, fill, spacing=7):
    x1, y1, x2, y2 = box
    max_width = x2 - x1 - 44
    lines = []
    current = ""
    for ch in text:
        trial = current + ch
        if draw.textlength(trial, font=fnt) > max_width and current:
            lines.append(current)
            current = ch
        else:
            current = trial
    if current:
        lines.append(current)
    heights = [draw.textbbox((0, 0), line, font=fnt)[3] for line in lines]
    total = sum(heights) + spacing * (len(lines) - 1)
    y = y1 + (y2 - y1 - total) / 2
    for line, h in zip(lines, heights):
        w = draw.textlength(line, font=fnt)
        draw.text((x1 + (x2 - x1 - w) / 2, y), line, font=fnt, fill=fill)
        y += h + spacing


def rounded_box(draw, box, fill, outline, radius=26, width=4):
    draw.rounded_rectangle(box, radius=radius, fill=fill, outline=outline, width=width)


def arrow(draw, start, end, color=TEAL, width=10):
    if not color.startswith("#"):
        color = "#" + color
    draw.line([start, end], fill=color, width=width)
    import math
    angle = math.atan2(end[1] - start[1], end[0] - start[0])
    length = 28
    wing = 0.58
    p1 = (end[0] - length * math.cos(angle - wing), end[1] - length * math.sin(angle - wing))
    p2 = (end[0] - length * math.cos(angle + wing), end[1] - length * math.sin(angle + wing))
    draw.polygon([end, p1, p2], fill=color)


def save_canvas(image, path):
    image.save(path, "PNG", optimize=True)


def create_care_loop(path):
    img = Image.new("RGB", (2100, 1180), "#F7FAFB")
    d = ImageDraw.Draw(img)
    d.text((1050, 70), "双向家庭关怀闭环", font=font(72, True), fill="#17324D", anchor="ma")
    d.text((1050, 140), "AI 负责理解与协同，真实家人完成关怀行动", font=font(34), fill="#667784", anchor="ma")
    boxes = [
        (110, 340, 490, 570, "老人自然表达", "#FFF3E0", "#D9923B"),
        (650, 250, 1050, 480, "AI 识别事件、需求与风险", "#DDEFEA", "#267C7A"),
        (1270, 340, 1660, 570, "老人确认分享范围", "#E8EEF5", "#2E74B5"),
        (1500, 760, 1910, 990, "子女收到可行动摘要", "#E8EEF5", "#2E74B5"),
        (850, 820, 1260, 1050, "子女留言、回拨或处理", "#FFF3E0", "#D9923B"),
        (190, 760, 600, 990, "真实反馈回到老人", "#DDEFEA", "#267C7A"),
    ]
    for x1, y1, x2, y2, txt, fill, outline in boxes:
        rounded_box(d, (x1, y1, x2, y2), fill, outline)
        centered_text(d, (x1, y1, x2, y2), txt, font(43, True), "#17324D")
    arrow(d, (490, 455), (650, 385))
    arrow(d, (1050, 385), (1270, 455))
    arrow(d, (1660, 540), (1740, 760))
    arrow(d, (1500, 880), (1260, 930))
    arrow(d, (850, 930), (600, 880))
    arrow(d, (300, 760), (280, 570))
    save_canvas(img, path)


def create_architecture(path):
    img = Image.new("RGB", (2200, 1420), "#F7FAFB")
    d = ImageDraw.Draw(img)
    d.text((1100, 60), "归音 MVP 系统架构", font=font(72, True), fill="#17324D", anchor="ma")
    d.text((1100, 135), "模块化单体 + 独立实时网关 + AI Worker", font=font(34), fill="#667784", anchor="ma")
    layers = [
        (220, 230, 1980, 420, "交互层", ["老人模式", "子女模式", "运营后台"], "#E8EEF5", "#2E74B5"),
        (220, 500, 1980, 690, "接入层", ["API Gateway", "WebSocket 实时网关", "鉴权 / 限流 / 风控"], "#DDEFEA", "#267C7A"),
        (220, 770, 1980, 1005, "业务与 AI 层", ["用户家庭服务", "记忆 / RAG", "摘要与任务", "AI Orchestrator"], "#FFF3E0", "#D9923B"),
        (220, 1085, 1980, 1280, "数据与基础设施", ["PostgreSQL + pgvector", "Redis", "对象存储", "监控与审计"], "#EEF2F5", "#607789"),
    ]
    for x1, y1, x2, y2, label, nodes, fill, outline in layers:
        rounded_box(d, (x1, y1, x2, y2), "#FFFFFF", outline, radius=32, width=5)
        d.rounded_rectangle((x1, y1, x1 + 245, y2), radius=29, fill=fill)
        centered_text(d, (x1, y1, x1 + 245, y2), label, font(38, True), "#17324D")
        available = x2 - (x1 + 290)
        gap = 24
        w = (available - gap * (len(nodes) - 1)) / len(nodes)
        for idx, node in enumerate(nodes):
            nx1 = x1 + 285 + idx * (w + gap)
            nx2 = nx1 + w
            rounded_box(d, (nx1, y1 + 43, nx2, y2 - 43), fill, outline, radius=22, width=3)
            centered_text(d, (nx1, y1 + 43, nx2, y2 - 43), node, font(31, True), "#17324D")
    for y in (420, 690, 1005):
        arrow(d, (1100, y + 6), (1100, y + 73), color="#267C7A", width=9)
    save_canvas(img, path)


def create_privacy(path):
    img = Image.new("RGB", (2100, 1250), "#F7FAFB")
    d = ImageDraw.Draw(img)
    d.text((1050, 70), "授权感知的数据安全边界", font=font(72, True), fill="#17324D", anchor="ma")
    d.text((1050, 145), "先判断权限，再检索、生成与共享", font=font(34), fill="#667784", anchor="ma")
    columns = [
        (95, 280, 480, 1020, "老人数据域", ["原始对话", "候选记忆", "私密事件"], "#FFF3E0", "#D9923B"),
        (585, 280, 970, 1020, "授权决策层", ["用途", "对象", "期限", "敏感等级"], "#DDEFEA", "#267C7A"),
        (1075, 280, 1460, 1020, "派生信息域", ["结构化事件", "主题趋势", "待确认需求"], "#E8EEF5", "#2E74B5"),
        (1565, 280, 1950, 1020, "家庭可见域", ["授权摘要", "已确认需求", "行动状态"], "#EEF2F5", "#607789"),
    ]
    for x1, y1, x2, y2, title, items, fill, outline in columns:
        rounded_box(d, (x1, y1, x2, y2), "#FFFFFF", outline, radius=30, width=5)
        d.rounded_rectangle((x1, y1, x2, y1 + 155), radius=27, fill=fill)
        centered_text(d, (x1, y1, x2, y1 + 155), title, font(39, True), "#17324D")
        y = y1 + 210
        for item in items:
            rounded_box(d, (x1 + 45, y, x2 - 45, y + 120), fill, outline, radius=18, width=2)
            centered_text(d, (x1 + 45, y, x2 - 45, y + 120), item, font(31), "#243746")
            y += 155
    for x in (480, 970, 1460):
        arrow(d, (x + 10, 650), (x + 95, 650), color="#267C7A", width=9)
    d.rounded_rectangle((400, 1090, 1700, 1190), radius=20, fill="#17324D")
    d.text((1050, 1140), "任何共享均可追溯、可纠正、可撤回、可删除", font=font(38, True), fill="#FFFFFF", anchor="mm")
    save_canvas(img, path)


def create_roadmap(path):
    img = Image.new("RGB", (2200, 1080), "#F7FAFB")
    d = ImageDraw.Draw(img)
    d.text((1100, 65), "12–16 周 MVP 研发路线", font=font(72, True), fill="#17324D", anchor="ma")
    d.text((1100, 140), "先验证需求与技术，再进入真实家庭试点", font=font(34), fill="#667784", anchor="ma")
    y = 545
    d.line((210, y, 1990, y), fill="#B9C8D2", width=14)
    stages = [
        (260, "01", "需求验证", "第 1–2 周", "访谈与原型", "#D9923B"),
        (670, "02", "技术验证", "第 3–4 周", "语音链路与选型", "#267C7A"),
        (1080, "03", "核心开发", "第 5–9 周", "家庭、RAG 与周报", "#2E74B5"),
        (1490, "04", "安全适老", "第 10–12 周", "风控、隐私与测试", "#607789"),
        (1900, "05", "封闭试点", "第 13–16 周", "5–10 户家庭", "#17324D"),
    ]
    for x, num, title, weeks, note, color in stages:
        d.ellipse((x - 62, y - 62, x + 62, y + 62), fill=color, outline="#FFFFFF", width=8)
        d.text((x, y), num, font=font(34, True), fill="#FFFFFF", anchor="mm")
        d.text((x, y - 145), title, font=font(37, True), fill="#17324D", anchor="mm")
        d.text((x, y - 95), weeks, font=font(28), fill="#667784", anchor="mm")
        rounded_box(d, (x - 165, y + 120, x + 165, y + 285), "#FFFFFF", color, radius=22, width=4)
        centered_text(d, (x - 165, y + 120, x + 165, y + 285), note, font(29), "#243746")
    d.text((1100, 985), "验收重点：真实联网、授权闭环、证据可追溯、风险边界可验证", font=font(34, True), fill="#267C7A", anchor="mm")
    save_canvas(img, path)


def prepare_assets():
    ASSET_DIR.mkdir(parents=True, exist_ok=True)
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    cover = ASSET_DIR / "guiyin-cover-illustration.png"
    shutil.copy2(COVER_SOURCE, cover)
    diagrams = {
        "care_loop": ASSET_DIR / "care-loop.png",
        "architecture": ASSET_DIR / "system-architecture.png",
        "privacy": ASSET_DIR / "privacy-boundary.png",
        "roadmap": ASSET_DIR / "mvp-roadmap.png",
    }
    create_care_loop(diagrams["care_loop"])
    create_architecture(diagrams["architecture"])
    create_privacy(diagrams["privacy"])
    create_roadmap(diagrams["roadmap"])
    return cover, diagrams


def configure_styles(document):
    section = document.sections[0]
    section.page_width = Inches(PAGE_WIDTH_IN)
    section.page_height = Inches(PAGE_HEIGHT_IN)
    section.top_margin = Inches(MARGIN_IN)
    section.right_margin = Inches(MARGIN_IN)
    section.bottom_margin = Inches(MARGIN_IN)
    section.left_margin = Inches(MARGIN_IN)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = document.styles["Normal"]
    normal.font.name = FONT_LATIN
    normal.font.size = Pt(11)
    normal.font.color.rgb = rgb(INK)
    normal._element.rPr.rFonts.set(qn("w:ascii"), FONT_LATIN)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), FONT_LATIN)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_EAST_ASIA)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.333

    specs = {
        "Heading 1": (16, BLUE, 18, 10),
        "Heading 2": (13, BLUE, 12, 6),
        "Heading 3": (12, DARK_BLUE, 8, 4),
    }
    for name, (size, color, before, after) in specs.items():
        style = document.styles[name]
        style.font.name = FONT_LATIN
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = rgb(color)
        style._element.rPr.rFonts.set(qn("w:ascii"), FONT_LATIN)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), FONT_LATIN)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_EAST_ASIA)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.keep_together = True

    title = document.styles["Title"]
    title.font.name = FONT_DISPLAY
    title.font.size = Pt(30)
    title.font.bold = True
    title.font.color.rgb = rgb(NAVY)
    title._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_DISPLAY)
    title.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(8)

    subtitle = document.styles["Subtitle"]
    subtitle.font.name = FONT_LATIN
    subtitle.font.size = Pt(14)
    subtitle.font.color.rgb = rgb(TEAL)
    subtitle._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_EAST_ASIA)
    subtitle.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(18)


def configure_header_footer(document):
    section = document.sections[0]
    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run("归音  ·  家庭情感协同 AI App")
    set_run_font(r, size=8.5, color=MUTED, bold=True)
    add_bottom_border(p, color="D7E0E6", size=6, space=4)

    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    r = p.add_run("研发基线 v1.0   ·   第 ")
    set_run_font(r, size=8.5, color=MUTED)
    add_field(p, "PAGE")
    r = p.add_run(" 页 / 共 ")
    set_run_font(r, size=8.5, color=MUTED)
    add_field(p, "NUMPAGES")
    r = p.add_run(" 页")
    set_run_font(r, size=8.5, color=MUTED)


def add_cover(document, cover_path):
    section = document.sections[0]
    section.different_first_page_header_footer = True
    first_header = section.first_page_header
    first_header.paragraphs[0].text = ""
    first_footer = section.first_page_footer
    first_footer.paragraphs[0].text = ""

    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(10)
    r = p.add_run("GUIYIN  /  PROJECT BLUEPRINT")
    set_run_font(r, size=9.5, color=AMBER, bold=True)

    p = document.add_paragraph(style="Title")
    p.add_run("归音")
    p = document.add_paragraph(style="Subtitle")
    p.add_run("面向老人与家庭成员的双向情感协同 AI App")

    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(18)
    r = p.add_run("完整项目开发文档  ·  产品 / 技术 / 安全 / 研发 / 商业")
    set_run_font(r, size=10.5, color=MUTED)

    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    inline = p.add_run().add_picture(str(cover_path), width=Inches(6.48))
    set_alt_text(inline, "归音项目封面插画", "老人和异地子女通过可信 AI 助手保持联系的概念插画")

    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run("研发基线草案  v1.0")
    set_run_font(r, size=11, color=NAVY, bold=True)
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run("2026 年 8 月 30 日  ·  中国大陆试点版")
    set_run_font(r, size=9.5, color=MUTED)
    document.add_page_break()


def add_front_matter(document):
    p = document.add_paragraph("执行摘要", style="Heading 1")
    p.paragraph_format.space_before = Pt(0)
    add_callout(
        document,
        "项目定位",
        "归音不是替代子女的“数字家人”，而是一个听得懂老人、提醒得了子女、促成真实家庭联系的情感协同智能体。",
        fill=TEAL_LIGHT,
        accent=TEAL,
    )
    add_body_paragraph(
        document,
        "本项目面向独居、空巢老人及其异地子女，提供语音优先的 AI 对话、家庭知识检索、授权式长期记忆、需求转达、生活状态摘要和家庭留言闭环。首版以独立 Android/iOS App 交付，通过互联网接入真实 ASR、LLM 与 TTS 服务。",
    )
    add_body_paragraph(
        document,
        "技术上采用 Flutter 双端客户端、FastAPI 模块化后端、PostgreSQL + pgvector、Redis、对象存储与可替换的 AI 供应商适配层。产品不提供疾病诊断、用药决策或未经授权的监控。",
    )
    facts = [
        ["决策项", "研发基线"],
        ["首版形态", "一个独立 App，包含老人模式与子女模式；另设内部 Web 管理后台"],
        ["核心闭环", "老人表达 → AI 理解 → 老人授权 → 子女行动 → 真实反馈回到老人"],
        ["MVP 周期", "4–6 人兼职研究生团队，预计 12–16 周"],
        ["首版边界", "非医疗器械，不做诊断、不做持续监听、不正式上线一比一声线克隆"],
        ["北极星指标", "每户家庭每周由系统促成的有效真实关怀次数"],
    ]
    add_markdown_table(document, facts)
    add_callout(
        document,
        "开发建议",
        "先选择一个试点地区完成 10 户家庭封闭测试，再依据真实留存、摘要准确率和关怀行动数据决定是否扩大投入。",
        fill=AMBER_LIGHT,
        accent=AMBER,
    )
    document.add_page_break()
    p = document.add_paragraph("目录", style="Heading 1")
    p.paragraph_format.space_before = Pt(0)
    add_toc_entries(document, TOC_ENTRIES[:15])
    document.add_page_break()
    p = document.add_paragraph("目录（续）", style="Heading 1")
    p.paragraph_format.space_before = Pt(0)
    add_toc_entries(document, TOC_ENTRIES[15:])
    document.add_page_break()


PARTS = {
    "1": ("第一篇", "产品定义与需求", "从项目目标、用户价值到首版产品边界"),
    "7": ("第二篇", "技术架构与 AI 能力", "从联网语音链路到数据模型与接口"),
    "12": ("第三篇", "信任、安全与合规", "让授权、隐私和风险处置成为底层能力"),
    "16": ("第四篇", "体验、质量与验收", "把适老体验转化为可验证的工程标准"),
    "19": ("第五篇", "交付、团队与商业", "从研发节奏到试点和创业比赛表达"),
    "25": ("第六篇", "工程落地与上线", "仓库、环境、检查清单与成功定义"),
}


def add_part_divider(document, number, title, subtitle):
    if number != "第一篇":
        document.add_page_break()
    p = document.add_paragraph()
    p.paragraph_format.space_before = Pt(100)
    p.paragraph_format.space_after = Pt(12)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(number)
    set_run_font(r, size=10, color=AMBER, bold=True)
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(14)
    r = p.add_run(title)
    set_run_font(r, east_asia=FONT_DISPLAY, size=25, color=NAVY, bold=True)
    add_bottom_border(p, color=TEAL, size=16, space=10)
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(12)
    r = p.add_run(subtitle)
    set_run_font(r, size=11.5, color=MUTED)
    document.add_page_break()


def add_figure(document, path, title, description):
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)
    shape = p.add_run().add_picture(str(path), width=Inches(6.35))
    set_alt_text(shape, title, description)
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(10)
    r = p.add_run(title)
    set_run_font(r, size=9.5, color=MUTED, italic=True)
    set_keep_with_next(p, False)


def parse_table(lines, start):
    rows = []
    i = start
    while i < len(lines) and lines[i].strip().startswith("|"):
        raw = lines[i].strip().strip("|")
        cells = [c.strip() for c in raw.split("|")]
        if not all(re.fullmatch(r":?-{3,}:?", c.replace(" ", "")) for c in cells):
            rows.append(cells)
        i += 1
    width = len(rows[0]) if rows else 0
    rows = [r + [""] * (width - len(r)) for r in rows if len(r) <= width]
    return rows, i


def add_code_block(document, code_lines):
    for idx, line in enumerate(code_lines):
        p = document.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.18)
        p.paragraph_format.right_indent = Inches(0.08)
        p.paragraph_format.space_before = Pt(0 if idx else 5)
        p.paragraph_format.space_after = Pt(0 if idx < len(code_lines) - 1 else 6)
        p.paragraph_format.line_spacing = 1.0
        p.paragraph_format.keep_with_next = idx < len(code_lines) - 1
        set_paragraph_shading(p, "F2F6F8")
        r = p.add_run(line if line else " ")
        set_run_font(r, name=FONT_MONO, east_asia=FONT_EAST_ASIA, size=8.3, color=NAVY)


def add_body_from_markdown(document, diagrams):
    text = SOURCE_MD.read_text(encoding="utf-8")
    lines = text.splitlines()
    start = next(i for i, line in enumerate(lines) if line.startswith("## 1."))
    lines = lines[start:]
    bullet_num, decimal_num = ensure_numbering(document)
    in_code = False
    code_lines = []
    i = 0
    while i < len(lines):
        raw = lines[i]
        stripped = raw.strip()
        if stripped.startswith("```"):
            if in_code:
                add_code_block(document, code_lines)
                code_lines = []
                in_code = False
            else:
                in_code = True
            i += 1
            continue
        if in_code:
            code_lines.append(raw)
            i += 1
            continue
        if not stripped or stripped == "---":
            i += 1
            continue
        if stripped.startswith("|"):
            rows, i = parse_table(lines, i)
            add_markdown_table(document, rows)
            continue
        h = re.match(r"^(#{2,4})\s+(.+)$", stripped)
        if h:
            level = len(h.group(1)) - 1
            title = h.group(2)
            top_match = re.match(r"(\d+)\.\s", title)
            if level == 1 and top_match and top_match.group(1) in PARTS:
                add_part_divider(document, *PARTS[top_match.group(1)])
            p = document.add_paragraph(title, style=f"Heading {min(level, 3)}")
            if level == 1:
                add_bottom_border(p, color="D7E0E6", size=5, space=4)
            if title.startswith("2. 项目概述"):
                add_figure(document, diagrams["care_loop"], "图 1  双向家庭关怀闭环", "老人表达、AI 理解、老人授权、子女行动及真实反馈形成的闭环")
            elif title.startswith("7. 系统总体架构"):
                add_figure(document, diagrams["architecture"], "图 2  归音 MVP 系统总体架构", "交互层、接入层、业务与 AI 层以及数据基础设施的分层架构")
            elif title.startswith("12. 权限模型"):
                add_figure(document, diagrams["privacy"], "图 3  授权感知的数据安全边界", "从老人原始数据到家庭可见信息的分层授权和派生过程")
            elif title.startswith("19. 研发计划"):
                add_figure(document, diagrams["roadmap"], "图 4  12–16 周 MVP 研发路线", "从需求验证、技术验证、核心开发、安全适老化到封闭试点的五阶段路线")
            i += 1
            continue
        if stripped.startswith(">"):
            add_callout(document, "说明", stripped.lstrip("> "), fill=LIGHT, accent=BLUE)
            i += 1
            continue
        m_bullet = re.match(r"^-\s+(.+)$", stripped)
        if m_bullet:
            p = document.add_paragraph()
            apply_num(p, bullet_num)
            p.paragraph_format.space_after = Pt(4)
            p.paragraph_format.line_spacing = 1.208
            add_inline(p, m_bullet.group(1))
            i += 1
            continue
        m_num = re.match(r"^(\d+)[.)]\s+(.+)$", stripped)
        if m_num:
            p = document.add_paragraph()
            item_num_id = create_number_instance(document, decimal_num, int(m_num.group(1)))
            apply_num(p, item_num_id)
            p.paragraph_format.space_after = Pt(4)
            p.paragraph_format.line_spacing = 1.208
            add_inline(p, m_num.group(2))
            i += 1
            continue
        add_body_paragraph(document, stripped)
        i += 1
    if code_lines:
        add_code_block(document, code_lines)


def add_end_page(document):
    document.add_page_break()
    p = document.add_paragraph()
    p.paragraph_format.space_before = Pt(110)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("归音")
    set_run_font(r, east_asia=FONT_DISPLAY, size=30, color=NAVY, bold=True)
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("不是复制一个家人，而是让真实家人更及时地出现。")
    set_run_font(r, size=15, color=TEAL, bold=True)
    p.paragraph_format.space_after = Pt(24)
    add_callout(document, "文档状态", "本文件为研发基线 v1.0。模型供应商、试点地区、成本单价与正式合规路径应在技术验证阶段更新。", fill=LIGHT, accent=BLUE)


def set_core_properties(document):
    props = document.core_properties
    props.title = "归音项目开发文档"
    props.subject = "面向老人与家庭成员的双向情感协同 AI App"
    props.author = "归音项目组"
    props.keywords = "适老化, 情感AI, RAG, 语音交互, 家庭关怀"
    props.comments = "研发基线 v1.0"


def main():
    cover, diagrams = prepare_assets()
    document = Document()
    configure_styles(document)
    configure_header_footer(document)
    set_update_fields(document)
    set_core_properties(document)
    add_cover(document, cover)
    add_front_matter(document)
    add_body_from_markdown(document, diagrams)
    add_end_page(document)
    document.save(OUTPUT_DOCX)
    print(OUTPUT_DOCX)


if __name__ == "__main__":
    main()
